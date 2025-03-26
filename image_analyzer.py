import os
import sys
import datetime
import signal
import argparse
import glob
import traceback
import logging
import time  # Import the time module

from docx import Document
from PIL import Image
import io
from PyPDF2 import PdfReader
import ollama
from typing import List, Tuple
from tqdm import tqdm  # Import tqdm for progress bars
import re  # Import the regular expression module

# ANSI escape codes for colors
LIGHT_GREEN = "\033[92m"
PINK = "\033[95m"
YELLOW = "\033[93m"
LIGHT_BLUE = "\033[94m"
RESET_COLOR = "\033[0m"

# Configure logging (without color codes for clean logs)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Utility Functions ---
def get_current_datetime():
    """Returns the current date and time as a formatted string."""
    now = datetime.datetime.now()
    return now.strftime("%Y-%m-%d %H:%M:%S")

def check_folder_exists(folder_path):
    """Checks if a folder exists."""
    return os.path.isdir(folder_path)

def count_images_in_folder(folder_path):
    """Counts the number of image files in a folder. Supports common image formats."""
    image_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp')  # Add more if needed
    image_count = 0
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(image_extensions):
            image_count += 1
    return image_count

def read_text_from_pdf(pdf_path: str) -> str:
    """Reads text from a PDF file."""
    try:
        with open(pdf_path, 'rb') as file:  # Open in binary read mode
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        logging.error(f"Error reading PDF '{pdf_path}': {e}")
        return None

def read_text_from_docx(docx_path: str) -> str:
    """Reads text from a docx file."""
    try:
        document = Document(docx_path)
        full_text = []
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        logging.error(f"Error reading DOCX '{docx_path}': {e}")
        return None

def get_file_contents(file_path: str) -> str:
    """Reads content from a file based on its extension."""
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == '.pdf':
        return read_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return read_text_from_docx(file_path)
    elif file_extension in ('.txt', '.md'): # Add more text-based extensions
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logging.error(f"Error reading text file '{file_path}': {e}")
            return None
    else:
        logging.warning(f"Unsupported file type: {file_extension}")
        return None

def analyze_image_with_ollama(image_path: str, ollama_model: str, image_type: str = "general") -> str:
    """Analyzes an image using the specified Ollama model."""
    try:
        with open(image_path, 'rb') as image_file:
            image_data = image_file.read()

        if image_type == "receipt":
            prompt = (
                "Extract key information (field and value pairs) from this receipt image.  Provide the output in field-value pairs, without any surrounding formatting characters like asterisks or bolding.  Prioritize extracting text."
            )
        elif image_type == "handwritten":
            prompt = (
                "Transcribe this handwritten note as accurately as possible. Focus on capturing individual words and letters as they are written, preserving the original style as best as possible. Describe it, do not convert to any tabular format. "
            )
        else:  # "general"
            prompt = (
                "Describe the contents of this image in detail, mirroring its structure and style as closely as possible. "
                "Do not attempt to force the information into a tabular format unless the image is explicitly a table. "
                "Preserve the original layout and formatting of the text. If it appears to be a handwritten note, represent the text as handwritten text. "
                "If it's a receipt or bank document, mirror its layout."
            )

        response = ollama.generate(
            model=ollama_model,
            prompt=prompt,
            images=[image_data],  # Pass the image data directly
        )

        analysis_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', response['response'])  # Remove bold markdown
        return analysis_text

    except Exception as e:
        logging.error(f"Error analyzing image '{image_path}': {e}")
        return f"Error analyzing image: {e}"

def determine_image_type(image_path: str, ollama_model: str) -> str:
    """Determines the type of the image (receipt, handwritten, or general)."""
    try:
        with open(image_path, 'rb') as image_file:
            image_data = image_file.read()

        prompt = "What type of image is this? Is it a receipt, a handwritten note, or something else? Answer with just one word: receipt, handwritten, or general."  # Simple prompt
        response = ollama.generate(
            model=ollama_model,
            prompt=prompt,
            images=[image_data],
        )
        image_type = response['response'].strip().lower()

        if image_type not in ("receipt", "handwritten", "general"):
            logging.warning(f"Unexpected image type detected: {image_type}.  Defaulting to general.")
            return "general"  # Handle cases where the LLM gives unexpected output

        return image_type
    except Exception as e:
        logging.error(f"Error determining image type: {e}")
        return "general"  # Default to "general" in case of errors

def create_mirror_document_from_image(image_path: str, ollama_model: str, output_path: str):
    """Creates a document that mirrors the content of an image, using Ollama to describe it."""
    try:
        image_type = determine_image_type(image_path, ollama_model)
        analysis_text = analyze_image_with_ollama(image_path, ollama_model, image_type)

        document = Document()
        document.add_heading('Image Mirror Document', level=1)
        document.add_heading('Image Description', level=2)

        analysis_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', analysis_text)  # Remove bold markdown
        document.add_paragraph(analysis_text)

        # Add the image itself (optional)
        # document.add_picture(image_path, width=Inches(6)) # requires 'from docx.shared import Inches'

        document.save(output_path)
        logging.info(f"Mirror document created at: {output_path}")
    except Exception as e:
        logging.error(f"Error creating mirror document for '{image_path}': {e}")

def create_docx_from_analysis(analysis_results: List[Tuple[str, str]], output_path: str):
    """Creates a DOCX document from analysis results.  `analysis_results` is a list of (filename, analysis_text) tuples."""
    document = Document()
    document.add_heading('Analysis Results', level=1)

    for filename, analysis_text in analysis_results:
        document.add_heading(filename, level=2)  # Filename as a subheading

        analysis_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', analysis_text)  # Remove bold markdown
        document.add_paragraph(analysis_text)

    try:
        document.save(output_path)
        logging.info(f"Document saved to: {output_path}")
    except Exception as e:
        logging.error(f"Error saving DOCX document to '{output_path}': {e}")


# --- Main Function ---
def main():
    # --- Argument Parsing ---
    parser = argparse.ArgumentParser(description="Image and Text Analyzer using Ollama.")
    parser.add_argument("-m", "--model", required=True, help="Ollama model to use.")
    args = parser.parse_args()

    ollama_model = args.model

    # --- Signal Handling ---
    def signal_handler(sig, frame):
        print(f"{RESET_COLOR}\nCtrl+C detected. Exiting...{RESET_COLOR}")
        sys.exit(0)

    signal.signal(signal.SIGINT, signal_handler)

    # --- Initialization ---
    print(f"{LIGHT_GREEN}Current Date and Time: {get_current_datetime()}{RESET_COLOR}")
    print(f"{PINK}Using Ollama model: {ollama_model}{RESET_COLOR}")

    # Read folder path from folder.txt
    try:
        with open("folder.txt", "r") as f:
            folder_path = f.readline().strip()
    except FileNotFoundError:
        print(f"{YELLOW}Error: folder.txt not found. Please create it with the folder path.{RESET_COLOR}")
        sys.exit(1)
    except Exception as e:
        print(f"{YELLOW}Error reading folder.txt: {e}{RESET_COLOR}")
        sys.exit(1)

    if not check_folder_exists(folder_path):
        print(f"{YELLOW}Error: Folder '{folder_path}' does not exist.{RESET_COLOR}")
        sys.exit(1)

    image_count = count_images_in_folder(folder_path)
    print(f"{YELLOW}Folder '{folder_path}' exists and contains {image_count} images.{RESET_COLOR}")


    # --- Main Loop ---
    while True:
        try:
            prompt = input(f"{LIGHT_BLUE}Analyzer: {RESET_COLOR}")

            if not prompt:  # Handle empty prompts
                continue

            if prompt.lower() == "exit":
                break  # Exit the loop gracefully

            # --- Analyze images and print to terminal ---
            if "analyze the images in the folder, and translate the content here with field and value of the field" in prompt.lower():
                image_files = [f for f in glob.glob(os.path.join(folder_path, '*')) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'))]
                if not image_files:
                    print(f"{LIGHT_BLUE}No images found in the folder.{RESET_COLOR}")
                else:
                    for image_path in tqdm(image_files, desc=f"{LIGHT_BLUE}Analyzing images for terminal output{RESET_COLOR}", unit=f"{LIGHT_BLUE}image{RESET_COLOR}"):  # Use tqdm here
                        image_type = determine_image_type(image_path, ollama_model)  # Determine image type
                        analysis_text = analyze_image_with_ollama(image_path, ollama_model, image_type)  # Analyze using the correct prompt
                        print(f"{LIGHT_BLUE}Analysis of {os.path.basename(image_path)} (Type: {image_type}):\n{analysis_text}{RESET_COLOR}")

            # --- Create Mirror Document from Image ---
            elif "create mirror document of the image" in prompt.lower():
                image_files = [f for f in glob.glob(os.path.join(folder_path, '*')) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'))]
                if not image_files:
                    print(f"{LIGHT_BLUE}No images found in the folder.{RESET_COLOR}")
                else:
                    for image_path in tqdm(image_files, desc=f"{LIGHT_BLUE}Creating mirror documents{RESET_COLOR}", unit=f"{LIGHT_BLUE}image{RESET_COLOR}"):  # Use tqdm here
                        image_name = os.path.splitext(os.path.basename(image_path))[0]
                        output_docx_path = os.path.join(folder_path, f"{image_name}_mirror.docx")
                        create_mirror_document_from_image(image_path, ollama_model, output_docx_path)

            # --- File Analysis and Document Creation ---
            elif "analyze files and create a document" in prompt.lower():
                analysis_results = [] # List of tuples: (filename, analysis_text)
                all_files = glob.glob(os.path.join(folder_path, '*')) # Get all files in the directory

                for file_path in tqdm(all_files, desc=f"{LIGHT_BLUE}Analyzing files{RESET_COLOR}", unit=f"{LIGHT_BLUE}file{RESET_COLOR}"): # Use tqdm here

                    filename = os.path.basename(file_path)
                    if os.path.isfile(file_path):
                        if os.path.isfile(file_path) and any(file_path.lower().endswith(ext) for ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp')):
                            image_type = determine_image_type(file_path, ollama_model)  # Determine image type
                            analysis_text = analyze_image_with_ollama(file_path, ollama_model, image_type)  # Analyze using the correct prompt
                        else:
                            content = get_file_contents(file_path)
                            if content:
                                # Pass content to Ollama for text analysis
                                try:
                                    response = ollama.generate(model=ollama_model, prompt=f"Analyze the following text: {content}")
                                    analysis_text = response['response']
                                except Exception as e:
                                    analysis_text = f"Error analyzing text from '{filename}': {e}"
                                    logging.error(analysis_text)

                            else:
                                analysis_text = f"Skipped - Unsupported file type: {filename}"
                        analysis_results.append((filename, analysis_text))

                if analysis_results:
                    output_docx_path = "analysis_document.docx"
                    create_docx_from_analysis(analysis_results, output_docx_path)
                    print(f"{LIGHT_BLUE}Analysis complete. Document created at: {output_docx_path}{RESET_COLOR}")
                else:
                    print(f"{LIGHT_BLUE}No analyzable files found in the folder.{RESET_COLOR}")

            # --- Normal Chat with the LLM ---
            else:
                try:
                    response = ollama.generate(model=ollama_model, prompt=prompt)
                    print(f"{LIGHT_BLUE}Ollama: {response['response']}{RESET_COLOR}")
                except Exception as e:
                    print(f"{LIGHT_BLUE}Error during Ollama interaction: {e}{RESET_COLOR}")
                    logging.error(traceback.format_exc()) # Log the full stack trace


        except Exception as e:
            print(f"{LIGHT_BLUE}An unexpected error occurred: {e}{RESET_COLOR}")
            logging.error(traceback.format_exc()) # Log the full stack trace



if __name__ == "__main__":
    main()